import os
import base64
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
MAX_PDF_PAGES = 50
MAX_TEXT_LENGTH = 12000


def process_file(path: str, mime: str) -> dict:
    file_size = os.path.getsize(path)
    if file_size > MAX_FILE_SIZE:
        raise ValueError("File vượt quá 20MB. Vui lòng nén hoặc chụp ảnh từng trang.")

    if mime == "application/pdf":
        return _process_pdf(path)
    elif mime in ("image/jpeg", "image/png", "image/jpg"):
        return _process_image(path, mime)
    elif "wordprocessingml" in mime or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _process_docx(path)
    else:
        raise ValueError(f"Định dạng file không được hỗ trợ: {mime}")


def _process_pdf(path: str) -> dict:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    pages = min(len(doc), MAX_PDF_PAGES)
    text = "\n".join(doc[i].get_text() for i in range(pages))
    doc.close()
    return {"type": "text", "content": text[:MAX_TEXT_LENGTH], "file_type": "pdf"}


def _process_image(path: str, mime: str) -> dict:
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    return {"type": "image", "data": b64, "media_type": mime, "file_type": "image"}


def _process_docx(path: str) -> dict:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    text = "\n".join(paragraphs)
    return {"type": "text", "content": text[:MAX_TEXT_LENGTH], "file_type": "docx"}


def build_claude_content(file_result: dict, user_message: str) -> list[dict]:
    if file_result["type"] == "image":
        return [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": file_result["media_type"],
                    "data": file_result["data"],
                },
            },
            {"type": "text", "text": "Đây là file ảnh của tôi. " + user_message},
        ]
    else:
        return [
            {
                "type": "document",
                "source": {
                    "type": "text",
                    "media_type": "text/plain",
                    "data": file_result["content"],
                },
            },
            {"type": "text", "text": user_message},
        ]


def build_openai_content(file_result: dict, user_message: str) -> list[dict]:
    if file_result["type"] == "image":
        mime = file_result["media_type"]
        b64 = file_result["data"]
        return [
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}"},
            },
            {"type": "text", "text": "Đây là file ảnh của tôi. " + user_message},
        ]
    else:
        # OpenAI không có document block — nhúng text thẳng vào message
        combined = f"[Nội dung tài liệu]\n{file_result['content']}\n\n{user_message}"
        return [{"type": "text", "text": combined}]


def build_content(file_result: dict, user_message: str, engine: str = "anthropic") -> list[dict]:
    """Return content blocks in the format required by the given AI engine."""
    if engine == "openai":
        return build_openai_content(file_result, user_message)
    return build_claude_content(file_result, user_message)


def cleanup_temp_file(path: str) -> None:
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception as e:
        logger.warning(f"Could not delete temp file {path}: {e}")
